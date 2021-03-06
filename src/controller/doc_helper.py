# -*- coding:utf-8 -*-
import os
import sys

import flask
import gc
from flask import request

from src import app
from src import db
from docx import Document
from src.controller.common_function import check_directory, check_file
from src.entity.diagnose import Diagnose
from src.entity.difficulty_assessment import Difficulty_assessment
from src.entity.illness_history import Illness_history
from src.entity.non_surgical import Non_surgical
from src.entity.oral_examination import Oral_examination
from src.entity.personal_history import Personal_history
from src.entity.risk_assessment import Risk_assessment
from src.entity.surgical import Surgical
from src.entity.tooth_location import Tooth_location
from src.entity.user import User
from src.entity.usphs import Usphs


@app.route('/medical-case-of-illness/doc',methods=['GET'])
def get_doc():
    tooth_id = (int)(request.args['tooth_id'])
    tooth_location = Tooth_location.query.filter_by(tooth_id=tooth_id).first()
    num_list = []
    step_dict ={'tooth_location':0,'illness_history':1,'oral_examination':2,'diagnose':3,'difficulty_assessment':4,'handle':5,}
    step_info_list=[]
    if tooth_location:
        step_string = tooth_location.step
        tooth_step_list = step_string.split(',')
        if '' in tooth_step_list:
            tooth_step_list.remove('')
        for i in range(len(tooth_step_list)):
            num_list.append((int)(tooth_step_list[i]))
        #for key,value in step_dict.items():
        #    for temp in num_list:
        #        if value==temp:
        #            step_info_list.append(key)
        user_id = tooth_location.user_id
        personal_history = Personal_history.query.filter_by(user_id= user_id).all()
        risk = Risk_assessment.query.filter_by(user_id=user_id).all()
        if personal_history:
            num_list.append(7)
        if risk :
            num_list.append(8)
    path = generate_doc(tooth_id,num_list)
    response = flask.Response(path)
    response.headers['Access-Control-Allow-Origin'] = '*'
    return response, 200

def generate_doc(tooth_id,step_info_list):
    document = Document()
    full_paras=[]
    manager_list = []
    for step in step_info_list:
        manager = doc_manager(tooth_id,step)
        manager_list.append(manager)
    for manager in manager_list:
        paras = manager.get_document()
        full_paras = full_paras+paras
        del manager
        gc.collect()
    for para in full_paras:
        if not '{[' in para.text:
            document.add_paragraph(para.text,para.style)
    check_directory(tooth_id)
    flag, path = check_file(tooth_id, (str)(tooth_id) + '.docx')
    if flag:
        os.remove(path)
    document.save(path)
    return path

class doc_manager:
    document = ''
    full_dict={}
    full_text = ''
    tooth_id = 0
    table =''
    def __init__(self,tooth_id,table):
        self.tooth_id=tooth_id
        self.table = table
        if table ==5:
            surgical = Surgical.query.filter_by(tooth_id=tooth_id).first()
            if not surgical:
                self.document = Document('./templete/handle2.docx')
            else:
                self.document = Document('./templete/handle1.docx')
        elif table == 8:
            self.document = Document('./templete/risk.docx')
        elif table == 0:
            self.document = Document('./templete/illness.docx')
        elif table ==1:
            self.document = Document('./templete/illness_history.docx')
        elif table == 7:
            self.document = Document('./templete/personal_history.docx')
        elif table == 2:
            self.document = Document('./templete/oral_examination.docx')
        elif table == 3:
            self.document = Document('./templete/diagnose.docx')
        elif table == 4:
            self.document = Document('./templete/difficulty_assessment.docx')
        elif table == 6:
            self.document = Document('./templete/usphs.docx')

    def get_document(self):
        paragraphs = self.document.paragraphs
        for paragraph in paragraphs:
            self.full_text = self.full_text+'%separator%'+paragraph.text
        if self.table ==5:
            self._get_handle_dict()
        elif self.table ==8:
            self._get_risk_dict()
        else :
            self._get_illness_dict()
        for key, value in self.full_dict.items():
            self.full_text = self.full_text.replace(key, str(self.full_dict[key]))
        text_list = self.full_text.split('%separator%')
        text_list.remove('')
        for i in range(len(text_list)):
            if len(paragraphs[i].runs) > 1:
                paragraphs[i].text = text_list[i]
        for i in range(len(paragraphs)):
            for j in range(len(paragraphs[i].runs)):
                paragraphs[i].runs[j].style=self.document.paragraphs[i].runs[j].style
        return paragraphs


    def _get_handle_dict(self):
        reload(sys)
        sys.setdefaultencoding('utf8')
        surgical = Surgical.query.filter_by(tooth_id =self.tooth_id).first()
        if not surgical:
            surgical = Non_surgical.query.filter_by(tooth_id=self.tooth_id).first()
        dict = surgical.get_dict()
        if dict['handle_type'] == 1:
            if not dict['appease_medicine'] == '':
                dict['appease_medicine'] = '安抚药物{0},'.format(dict['appease_medicine'])
                dict['observed_time'] = '观察时间{0}。'.format(dict['observed_time'])
            if not dict['modulo'] == '':
                dict['modulo'] = '取模材料{0}。'.format(dict['modulo'])
                dict['inlay'] = '嵌体材料{0}。'.format(dict['inlay'])
            if dict['specific_method'] == '牙安抚治疗&树脂充填修复':
                dict['is_first'] = ''
                dict['is_second'] = ''
            if dict['specific_method'] == '嵌体修复':
                dict['qumo'] = ''
            if dict['etching_type'] == '自酸蚀粘接系统':
                dict['full_etching'] = dict['self_etching']
            if dict['microscope'] == '是':
                dict['microscope'] = '使用显微镜'
            if dict['microscope'] == '否':
                dict['microscope'] = '不使用显微镜'
        elif dict['handle_type'] == 0:
            if dict['specific_method'] == '药物治疗':
                dict['non_surgical'] = '	药物治疗：将药物氟化物:{0}，硝酸银:{1}涂布于龋损处30s。' \
                    .format(dict['fluorination'], dict['silver_nitrate'])
            elif dict['specific_method'] == '再矿化治疗':
                dict['non_surgical'] = '	再矿化治疗：患牙清洁，干燥，将矿化液浸湿的小棉球置于患牙牙面，反复涂搽3-4次。'
            elif dict['specific_method'] == '窝沟封闭':
                dict['non_surgical'] = '	窝沟封闭：1.清洁牙面： 在低速手机上装好{0}，' \
                                            '蘸取适量{1}于牙面，对牙面和窝沟来回刷洗1分钟，' \
                                            '同时不断滴水保持毛刷湿润。2.用棉纱球隔湿,压缩空气牙面吹干，' \
                                            '{2}蘸取酸蚀剂置于牙尖斜面的2／3上。酸蚀时间{3}。' \
                                            ' 3.流水冲洗牙面10-15秒，去除牙釉质表面和反应沉淀物。' \
                                            '4.洗刷笔蘸取适量封闭剂沿窝沟从远中向近中涂布在酸蚀后的牙面上。' \
                                            '5.1-2分钟{4}灯离牙尖1mm照射20-40秒。' \
                                            '6. 探针进行检查，调合，定期{5}复查。'.format(
                    dict['additional_device'], dict['reagent'], dict['tools'],dict['time_of_etching'],
                    dict['lamp'], dict['check_time'],
                )

        for key, value in dict.items():
            self.full_dict['{[' + key + ']}'] = dict[key]
    def _get_risk_dict(self):
        reload(sys)
        sys.setdefaultencoding('utf8')
        tooth_lcoation  = Tooth_location.query.filter_by(tooth_id  = self.tooth_id).first()
        risk = Risk_assessment.query.filter_by(user_id = tooth_lcoation.user_id).first()
        user = User.query.filter_by(user_id = tooth_lcoation.user_id).first()
        dit = dict(risk.get_dict().items()+user.get_dict().items())
        for key, value in dit.items():
            if value == '是':
                dit[key] = '有'
            elif value == '否':
                dit[key] = '无'
        if dit['gender'] == True:
            dit['gender'] = '女'
        else:
            dit['gender'] = '男'
        for key, value in dit.items():
            self.full_dict['{[' + key + ']}'] = dit[key]
    def _get_illness_dict(self):
        reload(sys)
        sys.setdefaultencoding('utf8')
        dic_list =[]
        user_id=0
        oral_tooth_location=''
        tooth_location = db.session.query(Tooth_location).filter(Tooth_location.tooth_id==self.tooth_id).first()
        if tooth_location:
            dic_list.append(tooth_location)
            user_id = tooth_location.user_id
        user = User.query.filter_by(user_id=user_id).first()
        if user:
            dic_list.append(user)
        illness_history = Illness_history.query.filter_by(tooth_id=self.tooth_id).first()
        if illness_history:
            dic_list.append(illness_history)
        personal_history = Personal_history.query.filter_by(user_id=user_id).first()
        if personal_history:
            dic_list.append(personal_history)
        usphs=Usphs.query.filter_by(user_id=user_id).first()
        if usphs:
            dic_list.append(usphs)
        oral_examination = Oral_examination.query.filter_by(tooth_id=self.tooth_id).first()
        if oral_examination:
            dic_list.append(oral_examination)
            oral_tooth_location = oral_examination.tooth_location
            oral_examination_dict = oral_examination.__dict__
            del oral_examination_dict['tooth_location']
        diagnose = Diagnose.query.filter_by(tooth_id=self.tooth_id).first()
        if diagnose:
            dic_list.append(diagnose)
        difficulty_assessment = Difficulty_assessment.query.filter_by(tooth_id=self.tooth_id).first()
        if difficulty_assessment:
            dic_list.append(difficulty_assessment)
        full_dict = []
        for dit in dic_list:
            full_dict = full_dict+dit.__dict__.items()
        full_dict = dict(full_dict)
        full_dict['oral_tooth_location'] = oral_tooth_location
        full_dict['tooth_location'] = tooth_location.tooth_location
        if tooth_location:
            if full_dict['is_fill_tooth'] == 0:
                full_dict['tooth_info'] = '不要求补牙'+full_dict['tooth_location'] + full_dict['symptom'] + full_dict[
                    'time_of_occurrence']
            else:
                full_dict['tooth_info'] = '要求补牙'+full_dict['tooth_location'] + '要求补牙'
        if user:
            if full_dict['gender'] == True:
                full_dict['gender'] = '女'
            else:
                full_dict['gender'] = '男'
        if full_dict.has_key('difficulty_level'):
            if full_dict['difficulty_level']==1:
                full_dict['doctor_level']='III'
            elif full_dict['difficulty_level']==3:
                full_dict['doctor_level']='I'
            elif full_dict['difficulty_level']==2:
                full_dict['doctor_level'] = 'II'
        if full_dict.has_key('is_primary'):
            if full_dict['is_primary']==0:
                full_dict['first_illness']=''
            if full_dict['is_primary']==1:
                full_dict['second_illness']=''
        for key, value in full_dict.items():
            self.full_dict['{[' + key + ']}'] = full_dict[key]